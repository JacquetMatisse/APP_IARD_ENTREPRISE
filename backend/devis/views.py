from rest_framework import generics
from .models import Devis
from .serializers import DevisSerializer
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.http import FileResponse, Http404
from docx import Document
import os
from datetime import datetime
from docx2pdf import convert


class DevisListCreateView(generics.ListCreateAPIView):
    queryset = Devis.objects.all().order_by('-date_creation')
    serializer_class = DevisSerializer

class DevisWordView(APIView):
    def get(self, request, pk):
        try:
            devis = Devis.objects.get(pk=pk)
        except Devis.DoesNotExist:
            raise Http404

        # Génération du document Word
        doc = Document()
        doc.add_heading('TARIFICATION INDICATIVE', 0)
        doc.add_paragraph(f"Numéro d'opportunité : {devis.numero_opportunite}")
        doc.add_paragraph(f"Nom du client : {devis.nom_client}")
        doc.add_paragraph(f"Type de garantie : {devis.type_garantie}")
        doc.add_paragraph(f"Type de bien : {devis.type_bien}")
        doc.add_paragraph(f"Destination de l'ouvrage : {devis.destination_ouvrage}")
        doc.add_paragraph(f"Types de travaux : {devis.types_travaux}")
        doc.add_paragraph(f"Coût de l'ouvrage : {devis.cout_ouvrage} €")
        doc.add_paragraph(f"Présence d'existant : {'Oui' if devis.presence_existant else 'Non'}")
        doc.add_paragraph(f"Client VIP : {'Oui' if devis.client_vip else 'Non'}")
        doc.add_paragraph(f"Souhaite RCMO : {'Oui' if devis.souhaite_rcmo else 'Non'}")
        # Initialisation des variables de logique d'affichage
        garantie = devis.type_garantie
        isDO = garantie == "DO seule"
        isTRC = garantie == "TRC seule"
        isDuo = garantie == "Duo"
        # Tableau des montants de garanties
        doc.add_paragraph('Montants de garanties :')
        table_gar = doc.add_table(rows=1, cols=2)
        table_gar.style = 'Table Grid'
        hdr_cells = table_gar.rows[0].cells
        hdr_cells[0].text = 'Libellé'
        hdr_cells[1].text = 'Montant'
        garanties = [
            ("Dommage matériels à l'ouvrage", devis.garantie_ouvrage),
            ("Responsabilité civile", devis.garantie_rc),
            ("Maintenance-visite", devis.garantie_maintenance),
            ("Mesure conservatoire", devis.garantie_mesure_conservatoire),
        ]
        for label, value in garanties:
            row_cells = table_gar.add_row().cells
            row_cells[0].text = label
            if isDO:
                row_cells[1].text = "SANS"
            else:
                row_cells[1].text = str(value) if value not in (None, "") else "SANS"

        doc.add_paragraph("")  # Saut de ligne

        # Tableau des montants de franchise
        doc.add_paragraph('Montants de franchise :')
        table_fra = doc.add_table(rows=1, cols=2)
        table_fra.style = 'Table Grid'
        hdr_cells = table_fra.rows[0].cells
        hdr_cells[0].text = 'Libellé'
        hdr_cells[1].text = 'Montant'
        franchises = [
            ("Dommage subis par les ouvrages de bâtiments", devis.franchise_batiment),
            ("Catastrophes naturelles", devis.franchise_cat_nat),
            ("Responsabilité civile - Assuré maître d'ouvrage", devis.franchise_rc_maitre),
            ("Responsabilité civile - Assuré intervenants", devis.franchise_rc_intervenant),
            ("Maintenance-visite", devis.franchise_maintenance),
        ]
        for label, value in franchises:
            row_cells = table_fra.add_row().cells
            row_cells[0].text = label
            if isTRC and label in ["Dommage subis par les ouvrages de bâtiments", "Catastrophes naturelles"]:
                row_cells[1].text = "SANS"
            else:
                row_cells[1].text = str(value) if value not in (None, "") else "SANS"

        # Ajout de la date de simulation
        doc.add_paragraph("") # Ajout d'un saut de ligne pour espacer
        now_formatted = datetime.now().strftime("%d/%m/%Y")
        doc.add_paragraph(f"Date de simulation de tarif : {now_formatted}")

        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Proposition_commerciale_{devis.numero_opportunite}_{now}.docx"
        filepath = os.path.join("/tmp", filename)
        doc.save(filepath)

        return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=filename)
    
class DevisPDFView(APIView):
    def get(self, request, pk):
        try:
            devis = Devis.objects.get(pk=pk)
        except Devis.DoesNotExist:
            raise Http404

        # Génère d'abord le Word
        doc = Document()
        doc.add_heading('TARIFICATION INDICATIVE', 0)
        doc.add_paragraph(f"Numéro d'opportunité : {devis.numero_opportunite}")
        doc.add_paragraph(f"Nom du client : {devis.nom_client}")
        doc.add_paragraph(f"Type de garantie : {devis.type_garantie}")
        doc.add_paragraph(f"Type de bien : {devis.type_bien}")
        doc.add_paragraph(f"Destination de l'ouvrage : {devis.destination_ouvrage}")
        doc.add_paragraph(f"Types de travaux : {devis.types_travaux}")
        doc.add_paragraph(f"Coût de l'ouvrage : {devis.cout_ouvrage} €")
        doc.add_paragraph(f"Présence d'existant : {'Oui' if devis.presence_existant else 'Non'}")
        doc.add_paragraph(f"Client VIP : {'Oui' if devis.client_vip else 'Non'}")
        doc.add_paragraph(f"Souhaite RCMO : {'Oui' if devis.souhaite_rcmo else 'Non'}")
        # Initialisation des variables de logique d'affichage
        garantie = devis.type_garantie
        isDO = garantie == "DO seule"
        isTRC = garantie == "TRC seule"
        isDuo = garantie == "Duo"
        # Tableau des montants de garanties
        doc.add_paragraph('Montants de garanties :')
        table_gar = doc.add_table(rows=1, cols=2)
        table_gar.style = 'Table Grid'
        hdr_cells = table_gar.rows[0].cells
        hdr_cells[0].text = 'Libellé'
        hdr_cells[1].text = 'Montant'
        garanties = [
            ("Dommage matériels à l'ouvrage", devis.garantie_ouvrage),
            ("Responsabilité civile", devis.garantie_rc),
            ("Maintenance-visite", devis.garantie_maintenance),
            ("Mesure conservatoire", devis.garantie_mesure_conservatoire),
        ]
        for label, value in garanties:
            row_cells = table_gar.add_row().cells
            row_cells[0].text = label
            if isDO:
                row_cells[1].text = "SANS"
            else:
                row_cells[1].text = str(value) if value not in (None, "") else "SANS"

        doc.add_paragraph("")  # Saut de ligne

        # Tableau des montants de franchise
        doc.add_paragraph('Montants de franchise :')
        table_fra = doc.add_table(rows=1, cols=2)
        table_fra.style = 'Table Grid'
        hdr_cells = table_fra.rows[0].cells
        hdr_cells[0].text = 'Libellé'
        hdr_cells[1].text = 'Montant'
        franchises = [
            ("Dommage subis par les ouvrages de bâtiments", devis.franchise_batiment),
            ("Catastrophes naturelles", devis.franchise_cat_nat),
            ("Responsabilité civile - Assuré maître d'ouvrage", devis.franchise_rc_maitre),
            ("Responsabilité civile - Assuré intervenants", devis.franchise_rc_intervenant),
            ("Maintenance-visite", devis.franchise_maintenance),
        ]
        for label, value in franchises:
            row_cells = table_fra.add_row().cells
            row_cells[0].text = label
            if isTRC and label in ["Dommage subis par les ouvrages de bâtiments", "Catastrophes naturelles"]:
                row_cells[1].text = "SANS"
            else:
                row_cells[1].text = str(value) if value not in (None, "") else "SANS"

        # Ajout de la date de simulation
        doc.add_paragraph("") # Ajout d'un saut de ligne pour espacer
        now_formatted = datetime.now().strftime("%d/%m/%Y")
        doc.add_paragraph(f"Date de simulation de tarif : {now_formatted}")

        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_docx = f"Proposition_commerciale_{devis.numero_opportunite}_{now}.docx"
        filename_pdf = f"Proposition_commerciale_{devis.numero_opportunite}_{now}.pdf"
        filepath_docx = os.path.join("/tmp", filename_docx)
        filepath_pdf = os.path.join("/tmp", filename_pdf)
        doc.save(filepath_docx)

        # Conversion en PDF
        convert(filepath_docx, filepath_pdf)

        return FileResponse(open(filepath_pdf, 'rb'), as_attachment=True, filename=filename_pdf)